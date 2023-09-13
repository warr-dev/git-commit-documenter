# Import necessary modules
from git import Repo, InvalidGitRepositoryError
from datetime import datetime, timedelta
from docx import Document
import os


# define the date range
start_date = datetime(2023, 5, 16)
end_date = datetime(2023, 5, 17)

# repos path
path = "D:\Work\Github\\1bit"
folders = next(os.walk(path))[1]


def get_commits_per_project(repos):
    # Define the keywords used to classify the commits
    ignore_keywords = ['merge branch']
    try:
        repo = Repo(repos)
    except InvalidGitRepositoryError:
        print(repos+' is not a git repo, skipping')
        return
    # Open the Git repository
    # repo = Repo(repos)
    repo.git.fetch()
    # Create a dictionary to store the commit classifications by author
    commit_dict = {}

    # Iterate through each commit and classify it based on the commit message
    for commit in repo.iter_commits():
        # if start_date <= commit.committed_datetime <= end_date:
        if (True):
            message = commit.message.lower()
            if any(keyword in message for keyword in ignore_keywords):
                continue

            author = commit.author.name

            if author not in commit_dict:
                commit_dict[author] = {'project':,'commits': [], 'count': 0}
                # commit_dict[author][classification] += 1
            commit_dict[author]['count'] += 1
            commit_dict[author]['commits'].append(message)

    # repo_name = os.path.basename(repo.working_dir)

    # document = Document()
    # document.add_heading('Git Commit Report', 0)

    # for author, commits in commit_dict.items():
    #     document.add_paragraph(author, style='List Bullet')
    #     for messages in commits['commits']:
    #         for message in messages.split('\n')[:-1]:
    #             document.add_paragraph(message, style='List Bullet 2')

    # document.add_paragraph("\nGenerated on: " +
    #                        datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

    # # Save the document
    # out = 'output/'+repo_name+'.docx'
    # document.save(out)
    # print(f'Report saved to {out}')


for folder in folders:
    full_path = os.path.join(path, folder)
    # print(full_path)
    print_commits(full_path)


# Define the path to the Git repository
# repo_path = 'D:\Work\Fileserver'
# repo_path = 'D:\Work\laravel streaming'
