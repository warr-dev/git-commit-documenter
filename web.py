# app.py
from flask import Flask, request, render_template
from git import Repo, InvalidGitRepositoryError
from datetime import datetime, timedelta
from docx import Document
import os


app = Flask(__name__)

# repos path
path = "D:\Work\Github\\1bit"


def get_commits_per_project(repos, preset):
    # Define the keywords used to classify the commits
    ignore_keywords = ['merge branch']
    try:
        repo = Repo(repos)
    except InvalidGitRepositoryError:
        print(repos+' is not a git repo, skipping')
        return
    repo.git.fetch()
    # Create a dictionary to store the commit classifications by author
    commit_dict = {}
    repo_name = os.path.basename(repo.working_dir)
    print(repo_name)

    # Iterate through each commit and classify it based on the commit message
    for commit in repo.iter_commits(since=preset):
        # if start_date <= commit.committed_datetime <= end_date:
        if (True):
            message = commit.message.lower()
            print(message)
            if any(keyword in message for keyword in ignore_keywords):
                continue

            author = commit.author.name

            if author not in commit_dict:
                commit_dict[author] = {repo_name: {'commits': [], 'count': 0}}
                # commit_dict[author][classification] += 1
            commit_dict[author][repo_name]['count'] += 1
            commit_dict[author][repo_name]['commits'].append(message)
    return commit_dict


def printDocx(commits):
    document = Document()
    document.add_heading('Git Commit Report', 0)

    for author, projects in commits.items():
        document.add_paragraph(author, style='List Bullet')
        for project, values in projects.items():
            print(project)
            document.add_paragraph(project, style='List Bullet 2')
            for messages in values['commits']:
                for message in messages.split('\n')[:-1]:
                    document.add_paragraph(message, style='List Bullet 3')

    document.add_paragraph("\nGenerated on: " +
                           datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

    # Save the document
    out = 'output/commit_report.docx'
    document.save(out)
    print(f'Report saved to {out}')
    return out


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        folders = next(os.walk(path))[1]
        repos = []
        preset = request.form['preset']
        for folder in folders:
            full_path = os.path.join(path, folder)
            commits = get_commits_per_project(full_path, preset)
            repos.append(commits)

        merged_dict = {}
        for repo in repos:
            if repo is None:
                continue
            for author, projects in repo.items():
                if author not in merged_dict:
                    merged_dict[author] = projects
                else:
                    merged_dict[author].update(projects)
        out = printDocx(merged_dict)

        return render_template('download.html', url=out)
    return render_template('index.html', default_path=path)


if __name__ == '__main__':
    app.run(debug=True)
