# Import necessary modules
from git import Repo, InvalidGitRepositoryError
from datetime import datetime, timedelta
from docx import Document
import os


# define the date range
start_date = datetime(2023, 5, 16)
end_date = datetime(2023, 5, 17)

# repos path
path = "D:\Work\Github\\1bit"
folders = next(os.walk(path))[1]


def get_commits_per_project(repos):
    # Define the keywords used to classify the commits
    ignore_keywords = ['merge branch']
    try:
        repo = Repo(repos)
    except InvalidGitRepositoryError:
        print(repos+' is not a git repo, skipping')
        return
    repo.git.fetch()
    # Create a dictionary to store the commit classifications by author
    commit_dict = {}
    repo_name = os.path.basename(repo.working_dir)

    # Iterate through each commit and classify it based on the commit message
    for commit in repo.iter_commits(since='yesterday'):
        # if start_date <= commit.committed_datetime <= end_date:
        if (True):
            message = commit.message.lower()
            if any(keyword in message for keyword in ignore_keywords):
                continue

            author = commit.author.name

            if author not in commit_dict:
                commit_dict[author] = {repo_name: {'commits': [], 'count': 0}}
                # commit_dict[author][classification] += 1
            commit_dict[author][repo_name]['count'] += 1
            commit_dict[author][repo_name]['commits'].append(message)
    return commit_dict
def printDocx(commits):
    document = Document()
    document.add_heading('Git Commit Report', 0)

    for author, projects in commits.items():
        document.add_paragraph(author, style='List Bullet')
        for project, values in projects.items():
            document.add_paragraph(project, style='List Bullet 2')
            for messages in values['commits']:
                for message in messages.split('\n')[:-1]:
                    document.add_paragraph(message, style='List Bullet 3')

    document.add_paragraph("\nGenerated on: " +
                           datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

    # Save the document
    out = 'output.docx'
    document.save(out)
    print(f'Report saved to {out}')


repos = []
for folder in folders:
    full_path = os.path.join(path, folder)
    commits = get_commits_per_project(full_path)
    repos.append(commits)
# print(repos)
# merged_dict = {}
# for repo in repos:
#     if repo is None:
#         continue
#     for author, projects in repo.items():
#         if author not in merged_dict:
#             merged_dict[author] = projects
#         else:
#             merged_dict[author].update(projects)
# printDocx(merged_dict)
# for repo,ind in repos:
#     print(repo,ind)
# for author, project in repo:
#     for messages in commits['commits']:
#         for message in messages.split('\n')[:-1]:
#             document.add_paragraph(message, style='List Bullet 2')


# full_path = os.path.join(path, folders[0])
# commits = get_commits_per_project(full_path)
# print(commits)
